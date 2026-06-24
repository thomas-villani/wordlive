#!/usr/bin/env python
"""Generate the sample document used by the wordlive Tutorial.

Produces ``quarterly-report.docx`` next to this script — a short, realistic
status report with headings, body prose, and a budget table. The Tutorial
(docs/tutorial.md) drives wordlive against this document step by step.

The committed ``quarterly-report.docx`` is regenerated from this script, so the
binary stays reviewable: read *this* file to see exactly what's in the sample.

Run:
    uv run --with python-docx python examples/sample/build_quarterly_report.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document  # python-docx; only needed to (re)build the sample


def build() -> Document:
    doc = Document()

    doc.add_heading("Q2 Project Status Report", level=0)  # 'Title' style
    doc.add_paragraph("Prepared by the Project Office — for internal review.")

    # heading:3 — anchor ids share the paragraph index space (title=para:1,
    # subtitle=para:2), so the first Heading 1 is para:3 / heading:3.
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This report summarizes progress, risks, and budget for the second quarter.")
    doc.add_paragraph("We utilise a weekly status cadence to keep stakeholders aligned.")

    # heading:6
    doc.add_heading("Risks", level=1)
    doc.add_paragraph(
        "The single biggest risk this quarter is schedule slip on the integration milestone."
    )
    doc.add_paragraph("The vendor has promised the updated components as soon as possible.")
    doc.add_paragraph("Mitigation owners are assigned but not yet confirmed.")

    # heading:10
    doc.add_heading("Budget", level=1)
    doc.add_paragraph(
        "Spend remains within the approved envelope. The table below breaks it down by category."
    )
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Category", "Q1", "Q2"
    for category, q1, q2 in [
        ("Personnel", "$120,000", "$118,000"),
        ("Equipment", "$30,000", "$22,000"),
        ("Travel", "$8,000", "$6,500"),
    ]:
        row = table.add_row().cells
        row[0].text, row[1].text, row[2].text = category, q1, q2

    # heading:28
    doc.add_heading("Next Steps", level=1)
    doc.add_paragraph("Confirm mitigation owners.")
    doc.add_paragraph("Finalize the vendor delivery date.")
    doc.add_paragraph("Circulate this report for sign-off.")

    return doc


def main() -> None:
    out = Path(__file__).with_name("quarterly-report.docx")
    build().save(out)
    print(f"wrote {out}")


if __name__ == "__main__":
    main()
