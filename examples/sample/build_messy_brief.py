#!/usr/bin/env python
"""Generate the deliberately-messy sample used by the Linting guide.

Produces ``messy-brief.docx`` next to this script — a short project brief that
looks like something a human typed in a hurry: a heading with a hand-applied
font override, stray double spaces and trailing whitespace, a space before a
period, a left-aligned body at Word's default line spacing, and a budget table
whose numeric columns are left-aligned. Every blemish maps to a linter rule, so the guide
(docs/linting.md) can lint it, preview the fixes, and regularize it.

The committed ``messy-brief.docx`` is regenerated from this script, so the
binary stays reviewable: read *this* file to see exactly what's wrong with it.

Run:
    uv run --with python-docx python examples/sample/build_messy_brief.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document  # python-docx; only needed to (re)build the sample
from docx.oxml.ns import qn
from docx.shared import Pt


def _preserve(paragraph) -> None:
    """Mark a paragraph's runs ``xml:space="preserve"``.

    OOXML collapses leading / trailing / repeated spaces on load unless the run
    opts out, so python-docx would otherwise *normalize away* the very defects
    this sample exists to demonstrate — setting ``preserve`` keeps them literal.
    """
    for run in paragraph.runs:
        run._element.set(qn("xml:space"), "preserve")


def build() -> Document:
    doc = Document()

    doc.add_heading("Project Falcon — Weekly Brief", level=0)  # 'Title' style
    doc.add_paragraph("Prepared in a hurry, as these things are.")

    # heading:3 — a Heading 1 with a hand-applied font override (Arial 16pt on a
    # run whose style is something else): trips heading-font-consistent.
    h = doc.add_heading("Status", level=1)
    run = h.runs[0]
    run.font.name = "Arial"
    run.font.size = Pt(16)

    # Trailing whitespace + a space before the period — trailing-whitespace,
    # space-before-punctuation. (xml:space=preserve, or OOXML eats the spaces.)
    _preserve(doc.add_paragraph("Delivery is on track for the milestone review .   "))
    # A double space mid-sentence — double-space.
    _preserve(doc.add_paragraph("The  team closed the two blocking defects this week."))
    # Leading whitespace on a body paragraph — leading-whitespace.
    _preserve(doc.add_paragraph("  A follow-up on the vendor SLA is still pending."))

    # heading:7
    doc.add_heading("Budget", level=1)
    doc.add_paragraph("Spend by category is summarized below.")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Category", "Budget", "Used"
    for category, budget, used in [
        ("Personnel", "$120,000", "82%"),
        ("Equipment", "$30,000", "51%"),
        ("Travel", "$8,000", "44%"),
    ]:
        row = table.add_row().cells
        row[0].text, row[1].text, row[2].text = category, budget, used

    # heading:25
    doc.add_heading("Next Steps", level=1)
    doc.add_paragraph("Confirm the vendor SLA remedy.")
    doc.add_paragraph("Book the milestone review.")

    return doc


def main() -> None:
    out = Path(__file__).with_name("messy-brief.docx")
    build().save(out)
    print(f"wrote {out}")


if __name__ == "__main__":
    main()
